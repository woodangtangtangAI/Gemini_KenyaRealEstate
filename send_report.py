import sys
try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

import os
import json
import smtplib
import pandas as pd
import requests
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.application import MIMEApplication
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    pass
from dotenv import load_dotenv

# 1. 환경 변수 로드
base_path = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(base_path, ".env"))

GENAI_API_KEY = os.getenv("GEMINI_API_KEY")
SENDER_EMAIL = os.getenv("MY_EMAIL")
SENDER_PW = os.getenv("MY_PASSWORD")
RECEIVER_EMAIL = os.getenv("RECEIVER_EMAIL")

def set_cell_border(cell, **kwargs):
    """
    docx cell 테두리 스타일 지정용 헬퍼 함수
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def generate_and_send_report():
    print("🚀 AI 주간 리포트 '즉시 발행' 프로세스를 시작합니다...")
    
    base_path = os.path.dirname(os.path.abspath(__file__))
    results_dir = os.path.join(base_path, "분석 결과")
    os.makedirs(results_dir, exist_ok=True)
    
    # 2. 최신 마스터 데이터 로드 (분석 결과 폴더에서 찾기)
    file_list = [os.path.join(results_dir, f) for f in os.listdir(results_dir) if f.startswith("nairobi_master_data_") and f.endswith(".csv")]
    if not file_list:
        print("❌ 마스터 데이터를 찾을 수 없습니다.")
        return
        
    latest_file = max(file_list, key=os.path.getmtime)
    df = pd.read_csv(latest_file)
    latest_macro = df.iloc[-1]
    
    # 거시경제 데이터 안전하게 추출
    usd_rate = latest_macro.get('USD_KES_Rate', '데이터 없음') if 'USD_KES_Rate' in latest_macro.index else '데이터 없음'
    cbr_rate = latest_macro.get('CBR_Rate', '데이터 없음') if 'CBR_Rate' in latest_macro.index else '데이터 없음'
    remittance_val = latest_macro.get('Remittance_M_USD', '데이터 없음') if 'Remittance_M_USD' in latest_macro.index else '데이터 없음'
    
    # 거시경제 추세 분석 (분석 결과 폴더에서 엑셀 읽기)
    macro_trend_text = ""
    macro_xlsx = os.path.join(results_dir, "kenya_macro_history.xlsx")
    if os.path.exists(macro_xlsx):
        try:
            df_macro = pd.read_excel(macro_xlsx, engine='openpyxl')
            if len(df_macro) >= 2 and 'USD_KES_Rate' in df_macro.columns:
                recent = df_macro.tail(5)  # 최근 5주 데이터
                rates = recent['USD_KES_Rate'].dropna().tolist()
                dates = recent['Date'].astype(str).tolist()
                trend_pairs = [f"{d}: {r:.2f} KES" for d, r in zip(dates, rates)]
                macro_trend_text = f"\n    [최근 5주 환율 변동 추이]\n    " + " → ".join(trend_pairs)
                
                if len(rates) >= 2:
                    change = rates[-1] - rates[0]
                    direction = "상승(원화 대비 달러강세/KES약세)" if change > 0 else "하락(KES강세)"
                    macro_trend_text += f"\n    → 환율 흐름: {abs(change):.2f} KES {direction}"
        except Exception as e:
            print(f"  ⚠️ 거시경제 추세 분석 실패: {e}")
    
    # 머신러닝 상위 변수 가져오기 (분석 결과 폴더에서 찾기)
    top_features_path = os.path.join(results_dir, "top_features.txt")
    top_features_data = "추출된 변수 데이터 없음"
    if os.path.exists(top_features_path):
        try:
            with open(top_features_path, "r", encoding="utf-8") as f:
                top_features_data = f.read().strip()
        except Exception as e:
            print(f"  ⚠️ top_features.txt 읽기 실패: {e}")
            
    # 지역별 분석 데이터 가져오기 (분석 결과 폴더에서 찾기)
    regional_sales_text = ""
    regional_yield_text = ""
    regional_data = {}
    regional_path = os.path.join(results_dir, "regional_analysis.json")
    if os.path.exists(regional_path):
        try:
            with open(regional_path, 'r', encoding='utf-8') as f:
                regional_data = json.load(f)
            
            sales_lines = []
            for loc, stats in regional_data.get("sales", {}).items():
                sales_lines.append(
                    f"    - {loc}: 매물 {stats['count']}건, ㎡당 평균 {stats['avg_price_sqm']:,} KES (평균 매매가 {stats['avg_price']:,} KES)"
                )
            regional_sales_text = "\n    [나이로비 지역별 매매 가격 비교 (㎡당 단가 기준)]\n" + "\n".join(sales_lines)
            
            yield_lines = []
            for loc, y_val in regional_data.get("yields", {}).items():
                yield_lines.append(f"    - {loc}: 연간 임대 수익률 {y_val * 100:.2f}%")
            regional_yield_text = "\n    [지역별 매매가 대비 임대 수익률 (Yields)]\n" + "\n".join(yield_lines)
        except Exception as e:
            print(f"  ⚠️ 지역별 분석 데이터 로드 실패: {e}")
            
    # 최신 뉴스 가져오기 (분석 결과 폴더에서 찾기)
    news_text = ""
    news_path = os.path.join(results_dir, "scraped_news.json")
    if os.path.exists(news_path):
        try:
            with open(news_path, 'r', encoding='utf-8') as f:
                news_list = json.load(f)
            news_lines = []
            for item in news_list[:6]:
                news_lines.append(f"    - 기사 제목: {item['title']} (출처: {item['source']})")
            news_text = "\n    [케냐 현지 최신 부동산 및 개발 정책 관련 헤드라인 기사]\n" + "\n".join(news_lines)
        except Exception as e:
            print(f"  ⚠️ 뉴스 데이터 로드 실패: {e}")

    # 3. [개편된 3대 테마 프롬프트] 퀀트 애널리스트 버전
    prompt = f"""
    당신은 나이로비 부동산 시장을 전문적으로 정성적/정량적 데이터에 기반해 분석하는 퀀트 애널리스트(Quant Analyst)입니다.
    아래의 정량적 매크로 지표, 현지 실시간 뉴스 헤드라인, ㎡당 단가 및 임대수익률 통계를 종합 분석하여 주간 브리핑 보고서를 작성하세요.

    [테마 1: 거시경제 정량 지표]
    - 현재 환율: {usd_rate} KES/USD
    - 케냐 기준금리(CBR): {cbr_rate}%
    - 케냐 해외 송금액(Remittance): {remittance_val} Million USD (최신 월 통계)
    {macro_trend_text}

    [테마 2: 시장 수급 및 임대 수익률 지표]
    - 전체 분석 매물 수: {len(df):,}건
    {regional_sales_text}
    {regional_yield_text}

    [테마 3: 실시간 현지 부동산 뉴스 및 정책 동향 (정성적 지표)]
    {news_text}

    [모델 기반 변수 영향도]
    - XGBoost 모델 요인 영향도 상위 5개: {top_features_data}

    [보고서 서식 및 분석 지시사항]
    1. 위의 숫자를 대충 나열하지 말고, **숫자 간의 유기적 관계**를 설명하세요. (예: 해외 송금액 증가가 주택 구매 수요에 미친 영향, KES 환율 변동이 자재 수입 및 주택 가격에 미친 영향)
    2. 뉴스 헤드라인(예: 고층 아파트 공급 과잉설, 인프라 개발 소식 등)과 실제 수집된 특정 지역(Westlands, Kilimani 등)의 ㎡당 매매가 및 임대수익률 데이터를 매칭해 정성적 해석을 가미하세요.
    3. 금리 변동이 모기지 대출 및 실수요에 미칠 영향을 반영하여, 투자 관점의 시사점을 도출하세요.
    4. 보고서는 반드시 전문적인 톤앤매너로 서론-본론-결론 형식으로 작성하며 다음 목차를 철저히 지키십시오:
       - **요약 (Executive Summary)**: 이번 주의 핵심 이슈와 요약
       - **거시경제 동향**: 환율, 금리, 해외 송금액 등 금융 요소가 미친 영향
       - **나이로비 지역별 가격 및 수익률 동향**: 주요 Suburb의 ㎡당 가격 추이와 렌트 수익률 해석
       - **현지 시장 이슈 및 규제 뉴스 해설**: 수집된 뉴스 헤드라인 기사의 의미와 시장 파장 분석
       - **AI 및 머신러닝 예측 요인 분석**: XGBoost 피처 중요도가 시사하는 점
       - **투자자 제언 (Strategic Recommendations)**: 투자 의견(매수/매도/보유) 및 투자 시 주의해야 할 잠재적 리스크 요인
    """
    
    # 4. 내 API 키로 사용 가능한 모델 자동 탐색
    print("🔍 사용 가능한 AI 모델을 검색 중입니다...")
    target_model = "models/gemini-2.5-flash"
    available_models = []
    try:
        models_url = f"https://generativelanguage.googleapis.com/v1beta/models?key={GENAI_API_KEY}"
        model_res = requests.get(models_url, timeout=10).json()
        available_models = [m['name'] for m in model_res.get('models', []) if 'generateContent' in m.get('supportedGenerationMethods', [])]
        
        preferred = ["models/gemini-2.5-flash", "models/gemini-1.5-pro", "models/gemini-1.5-pro-latest"]
        for pm in preferred:
            if pm in available_models:
                target_model = pm
                break
        else:
            flash_models = [m for m in available_models if 'flash' in m]
            if flash_models:
                target_model = flash_models[0]
            
        print(f"✅ 모델 탐색 완료: {target_model}")
    except Exception as e:
        print(f"⚠️ 모델 검색 중 에러 발생 (기본값 설정): {e}")

    # 5. 선택된 모델과 통신 (API 키 유출 예외 처리 적용)
    print(f"🧠 {target_model} 모델에 3대 테마 기반 정밀 분석을 요청합니다...")
    ai_insight = ""
    is_key_leaked = False
    
    url = f"https://generativelanguage.googleapis.com/v1beta/{target_model}:generateContent?key={GENAI_API_KEY}"
    headers = {'Content-Type': 'application/json'}
    data = {
        "contents": [{"parts": [{"text": prompt}]}]
    }
    
    try:
        response = requests.post(url, headers=headers, json=data, timeout=120)
        response_data = response.json()
        
        if response.status_code == 403 and "leaked" in str(response_data).lower():
            is_key_leaked = True
            raise Exception("Gemini API key가 유출로 인해 차단된 상태입니다.")
        elif response.status_code != 200:
            raise Exception(f"API Error Code {response.status_code}: {response_data}")
            
        ai_insight = response_data['candidates'][0]['content']['parts'][0]['text']
        print("✅ 퀀트 분석 보고서 초안이 완벽히 도출되었습니다!")
    except Exception as e:
        print(f"❌ [AI 통신 실패] {e}")
        if is_key_leaked:
            print("\n" + "="*80)
            print("🚨 [경고] 사용 중인 GEMINI_API_KEY가 공개 깃허브 등에 노출되어 구글에 의해 무효화되었습니다.")
            print("💡 해결 방법:")
            print("  1. Google AI Studio (https://aistudio.google.com/)에서 새 API 키를 무료로 생성하세요.")
            print("  2. 로컬의 '.env' 파일 내 GEMINI_API_KEY 값을 새 키로 변경해 주세요.")
            print("  3. 깃허브 클라우드 연동을 사용 중이시라면 GitHub Repository -> Settings -> Secrets and Variables -> Actions 에 가셔서 GEMINI_API_KEY 값을 새 키로 업데이트해 주세요.")
            print("="*80 + "\n")
            
            ai_insight = (
                "## [시스템 공지 - AI 분석 일시 장애]\n\n"
                "현재 설정된 **Gemini API Key**가 깃허브 공개 저장소 노출 등의 이유로 **유출 차단(Leaked & Revoked)**되었습니다.\n"
                "그에 따라 AI 기반 마켓 센티먼트 및 전망 보고서의 동적 작성이 차단된 상태입니다.\n\n"
                "### 🛠️ 장애 해결 방법:\n"
                "1. **Google AI Studio** (https://aistudio.google.com/)에 접속하여 새 API 키를 재발급(무료)받습니다.\n"
                "2. **로컬 환경**: `G:\\내 드라이브\\[Python]\\케냐 부동산\\.env` 파일의 `GEMINI_API_KEY` 값을 새로 발급받은 키로 교체해 줍니다.\n"
                "3. **깃허브 클라우드 환경**: 깃허브 저장소(Gemini_KenyaRealEstate)의 `Settings -> Secrets and Variables -> Actions` 메뉴에서 `GEMINI_API_KEY` 값을 새 키로 업데이트해 줍니다."
            )
        else:
            ai_insight = f"[AI 정밀 분석 장애] 일시적 장애가 발생했습니다.\n상세 에러 내용: {e}"

    # 5.5 Word 파일 생성 및 저장 (분석 결과 폴더에 생성)
    current_date = datetime.now().strftime('%Y%m%d')
    docx_filename = f"나이로비_부동산_주간_브리핑_{current_date}.docx"
    docx_path = os.path.join(results_dir, docx_filename)
    
    print(f"📝 워드 문서 생성을 시작합니다: {docx_filename}")
    try:
        doc = Document()
        
        # 문서 기본 스타일 지정
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'
        font.size = Pt(10.5)
        
        # 제목
        title = doc.add_paragraph()
        title_run = title.add_run(f"나이로비 부동산 퀀트 분석 브리핑 ({datetime.now().strftime('%Y-%m-%d')})")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 요약 정보 테이블
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '총 수집 매물수'
        hdr_cells[1].text = '실시간 환율 (KES)'
        hdr_cells[2].text = '케냐 기준금리'
        hdr_cells[3].text = '송금 유입액 (USD)'
        
        # 헤더 셀 스타일 적용
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "004A99", "space": "0"})
            
        row_cells = table.add_row().cells
        row_cells[0].text = f"{len(df):,}건"
        row_cells[1].text = f"{usd_rate:.2f}" if isinstance(usd_rate, float) else str(usd_rate)
        row_cells[2].text = f"{cbr_rate}%"
        row_cells[3].text = f"{remittance_val:.2f}M" if isinstance(remittance_val, float) else str(remittance_val)
        
        doc.add_paragraph().paragraph_format.space_before = Pt(10)
        
        # 1. 거시경제 동향
        doc.add_heading('1. 거시경제 동향 및 시장 금융 지표', level=1)
        p_macro = doc.add_paragraph()
        p_macro.add_run("현재 케냐 부동산 시장을 둘러싼 거시경제 지표 변동 현황입니다.\n")
        p_macro.add_run(f"• USD/KES 환율: {usd_rate} KES/USD\n")
        p_macro.add_run(f"• 케냐 중앙은행 기준금리(CBR): {cbr_rate}%\n")
        p_macro.add_run(f"• Diaspora 해외송금 유동성: {remittance_val} Million USD\n")
        if macro_trend_text:
            p_macro.add_run(f"• 환율 변동 추이: {macro_trend_text.strip()}\n")
            
        # 2. 지역별 시세 및 수익률 현황 테이블 추가
        doc.add_heading('2. 지역별 매매 시세 및 임대 수익률 현황', level=1)
        
        sales_data = regional_data.get("sales", {})
        yields_data = regional_data.get("yields", {})
        
        if sales_data:
            t_reg = doc.add_table(rows=1, cols=4)
            t_reg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            th = t_reg.rows[0].cells
            th[0].text = '지역명'
            th[1].text = '㎡당 평균가'
            th[2].text = '평균 매매 가격'
            th[3].text = '연간 임대수익률'
            
            for cell in th:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "004A99", "space": "0"})
                
            for loc, stats in list(sales_data.items())[:12]:
                row = t_reg.add_row().cells
                row[0].text = loc
                row[1].text = f"{stats['avg_price_sqm']:,} KES"
                row[2].text = f"{stats['avg_price']:,} KES"
                
                y_val = yields_data.get(loc)
                row[3].text = f"{y_val * 100:.2f}%" if y_val else "N/A"
                
                for cell in row:
                    set_cell_border(cell, bottom={"sz": 4, "val": "single", "color": "E0E0E0", "space": "0"})
        else:
            doc.add_paragraph("데이터 분석 결과가 존재하지 않습니다.")
            
        # 3. AI 분석 및 인사이트 요약
        doc.add_heading('3. AI 퀀트 애널리스트 정밀 분석 리포트', level=1)
        for line in ai_insight.split('\n'):
            line_str = line.strip()
            if not line_str:
                continue
            if line_str.startswith('###'):
                doc.add_heading(line_str.replace('###', '').strip(), level=3)
            elif line_str.startswith('##'):
                doc.add_heading(line_str.replace('##', '').strip(), level=2)
            elif line_str.startswith('#'):
                doc.add_heading(line_str.replace('#', '').strip(), level=1)
            else:
                doc.add_paragraph(line_str)
                
        # 4. 모델 중요 피처
        doc.add_heading('4. 예측 모델의 중요 변수 기여도 (XGBoost)', level=1)
        doc.add_paragraph(f"XGBoost 머신러닝 모델의 집값 및 가격/sqm 요인 기여도 순위: {top_features_data}")
        
        doc.save(docx_path)
        print(f"✅ 워드 리포트 생성 및 저장 완료: {docx_path}")
    except Exception as e:
        print(f"⚠️ 워드 리포트 생성 실패: {e}")

    print("📧 이메일 발송 작업을 개시합니다...")
    
    # 6. 이메일 템플릿 작성
    current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    # 지역별 통계 테이블 HTML 생성
    regional_html = ""
    sales_data = regional_data.get("sales", {})
    yields_data = regional_data.get("yields", {})
    
    if sales_data:
        try:
            regional_html = '<table style="width:100%; border-collapse:collapse; margin:15px 0;">'
            regional_html += '<tr style="background:#004a99; color:white;"><th style="padding:8px; border:1px solid #ddd;">지역</th><th style="padding:8px; border:1px solid #ddd;">매물수</th><th style="padding:8px; border:1px solid #ddd;">㎡당 평균가 (KES)</th><th style="padding:8px; border:1px solid #ddd;">평균 가격 (KES)</th><th style="padding:8px; border:1px solid #ddd;">임대수익률</th></tr>'
            for i, (loc, stats) in enumerate(sales_data.items()):
                bg = '#f9f9f9' if i % 2 == 0 else '#ffffff'
                y_val = yields_data.get(loc)
                yield_str = f"{y_val * 100:.2f}%" if y_val else "데이터 부족"
                regional_html += f'<tr style="background:{bg};"><td style="padding:6px; border:1px solid #ddd; font-weight:bold;">{loc}</td><td style="padding:6px; border:1px solid #ddd; text-align:center;">{stats["count"]}</td><td style="padding:6px; border:1px solid #ddd; text-align:right;">{stats["avg_price_sqm"]:,}</td><td style="padding:6px; border:1px solid #ddd; text-align:right;">{stats["avg_price"]:,}</td><td style="padding:6px; border:1px solid #ddd; text-align:center; color:#d9534f; font-weight:bold;">{yield_str}</td></tr>'
            regional_html += '</table>'
        except Exception as e:
            print(f"  ⚠️ HTML 테이블 생성 에러: {e}")
            regional_html = ""
            
    # 최신 뉴스 기사 HTML 목록 생성
    news_html = ""
    if os.path.exists(news_path):
        try:
            with open(news_path, 'r', encoding='utf-8') as f:
                news_list = json.load(f)
            news_items_html = []
            for item in news_list[:6]:
                news_items_html.append(f"<li><a href='{item['link']}' style='color:#004a99; text-decoration:none; font-weight:bold;'>{item['title']}</a> <span style='color:#777; font-size:12px;'>({item['source']})</span></li>")
            news_html = "<ul style='padding-left:20px; line-height:1.8;'>" + "\n".join(news_items_html) + "</ul>"
        except:
            news_html = "<p>수집된 최신 뉴스가 없습니다.</p>"
            
    # 에러 안내 배너 (API 키 정지 시 배너 출력)
    warning_banner = ""
    if is_key_leaked:
        warning_banner = """
        <div style="background-color: #f2dede; border: 1px solid #ebccd1; color: #a94442; padding: 15px; border-radius: 4px; margin-bottom: 20px; font-family: sans-serif;">
            <strong style="font-size: 16px;">🚨 [긴급 공지] AI 분석 API Key 장애 안내</strong>
            <p style="margin: 5px 0 0 0; font-size: 14px;">현재 저장소의 <code>GEMINI_API_KEY</code>가 유출 차단 상태입니다. 이 보고서의 본문 분석 및 전망은 자동 복구 가이드로 대체되었습니다. 본래의 애널리스트 종합 보고서를 정상적으로 받아보시려면 수집 컴퓨터의 .env 및 깃허브 Action Secrets에 새 API 키를 재등록해 주셔야 합니다.</p>
        </div>
        """

    html_body = f"""
    <html>
    <body style="font-family: 'Malgun Gothic', sans-serif; color: #333; line-height: 1.6;">
        {warning_banner}
        <h2 style="color: #004a99; border-bottom: 2px solid #004a99; padding-bottom: 5px;">📊 나이로비 부동산 시장 전면 개편 퀀트 리포트</h2>
        <p><strong>보고서 발행일시:</strong> {current_time}</p>
        
        <div style="background: #f4f6f9; padding: 15px; border-radius: 5px; margin-bottom: 20px; border-left: 5px solid #004a99;">
            <h4 style="margin: 0 0 5px 0; color: #004a99;">📌 금주 주요 거시경제 지표</h4>
            <p style="margin: 0; font-size: 14px;">
                <strong>환율:</strong> {usd_rate:.2f} KES/USD | 
                <strong>기준금리(CBR):</strong> {cbr_rate}% | 
                <strong>Diaspora 해외 송금액:</strong> {remittance_val} Million USD
            </p>
        </div>
        
        <h3 style="margin-bottom: 10px; color:#004a99;">📍 Suburbs 시세 및 임대수익률 (Yields)</h3>
        {regional_html}
        
        <h3 style="margin-bottom: 10px; color:#004a99;">📰 케냐 현지 부동산 & 규제 정책 뉴스</h3>
        {news_html}
        
        <h3 style="margin-bottom: 10px; color:#004a99;">🧠 퀀트 애널리스트 AI Insight</h3>
        <div style="white-space: pre-wrap; background: #ffffff; border-left: 4px solid #004a99; padding: 15px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); font-size: 14px;">{ai_insight}</div>
        
        <h3 style="margin-top: 30px; color:#004a99;">🔍 주요 변수 영향도 (XGBoost ㎡당 단가 모델)</h3>
        <p style="font-size: 14px; background:#f9f9f9; padding:10px; border-radius:3px;">{top_features_data}</p>
        
        <p style="font-size: 12px; color:#999; margin-top:30px; text-align:center; border-top:1px solid #eee; padding-top:10px;">본 리포트는 수집된 데이터에 기반해 인공지능이 자동 작성한 시장 분석 보고서입니다.</p>
    </body>
    </html>
    """

    msg = MIMEMultipart()
    msg['Subject'] = f"[{datetime.now().strftime('%m/%d %H:%M:%S')}] 나이로비 부동산 3대 테마 퀀트 리포트"
    msg['From'] = SENDER_EMAIL
    
    recipients = [RECEIVER_EMAIL, "donghyun1.kwon@lge.com"]
    msg['To'] = ", ".join([r for r in recipients if r])
    
    msg.attach(MIMEText(html_body, 'html'))

    # 7. 이미지 첨부 (변수 중요도 + 지역별 비교) - 분석 결과 폴더에서 로드
    for img_name in ["feature_importance.png", "regional_price_comparison.png"]:
        img_path = os.path.join(results_dir, img_name)
        if os.path.exists(img_path):
            with open(img_path, 'rb') as f:
                msg.attach(MIMEImage(f.read(), name=img_name))
            
    # Word 파일 첨부
    if 'docx_path' in locals() and os.path.exists(docx_path):
        with open(docx_path, 'rb') as f:
            part = MIMEApplication(f.read(), Name=docx_filename)
            part['Content-Disposition'] = f'attachment; filename="{docx_filename}"'
            msg.attach(part)

    # 8. 이메일 발송
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(SENDER_EMAIL, SENDER_PW)
            server.send_message(msg, from_addr=SENDER_EMAIL, to_addrs=recipients)
        print("✅ [이메일 성공] 최신 리포트와 첨부파일이 지정된 수신처로 전송 완료되었습니다.")
    except Exception as e:
        print(f"❌ [이메일 실패] 메일 발송 중 오류: {e}")
        
    try:
        input("\n엔터 키를 누르면 창이 닫힙니다...")
    except EOFError:
        pass

if __name__ == "__main__":
    generate_and_send_report()